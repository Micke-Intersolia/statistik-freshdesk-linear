from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

doc.add_heading("Freshdesk Field Review — Bronze Layer", 0)
doc.add_paragraph(
    "Field inventory from freshdesk_snapshot_20260604T075839Z.json (5 091 tickets). "
    "Rec column: checkmark = include in agreed nightly snapshot field set, "
    "empty box = proposed to exclude."
)


def add_table(doc, heading, rows):
    doc.add_heading(heading, level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Field", "Rec", "Notes"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for field, rec, notes in rows:
        c = table.add_row().cells
        c[0].text = field
        c[1].text = rec
        c[2].text = notes
    doc.add_paragraph()


add_table(doc, "Core ticket fields (always present)", [
    ("id",         "✅", "Primary key"),
    ("subject",    "✅", "Ticket title"),
    ("status",     "✅", "Integer: 2=Open 3=Pending 4=Resolved 5=Closed"),
    ("priority",   "✅", "Integer: 1=Low 2=Medium 3=High 4=Urgent"),
    ("source",     "✅", "Integer: 1=Email 2=Portal 3=Phone etc."),
    ("type",       "✅", "199/200 non-null"),
    ("created_at", "✅", ""),
    ("updated_at", "✅", ""),
    ("due_by",     "✅", "191/200 non-null"),
    ("fr_due_by",  "✅", "First response deadline"),
    ("tags",       "✅", "Array — useful for categorisation"),
])

add_table(doc, "People & org fields", [
    ("requester_id",           "✅", "FK"),
    ("requester.id",           "⬜", "Same as requester_id"),
    ("requester.name",         "✅", "Customer name"),
    ("requester.email",        "✅", ""),
    ("requester.phone",        "⬜", "0/200 non-null in sample"),
    ("requester.mobile",       "⬜", "0/200 non-null in sample"),
    ("requester.contact_type", "⬜", "Low reporting value"),
    ("requester.first_seen",   "⬜", "Always null in sample"),
    ("requester.last_seen",    "⬜", "Always null in sample"),
    ("responder_id",           "✅", "Agent ID — 124/200 non-null"),
    ("company_id",             "✅", "FK"),
    ("company.name",           "✅", "141/200 — saves a lookup in silver"),
    ("group_id",               "✅", "Support group"),
    ("product_id",             "✅", "Which product / portal"),
    ("email_config_id",        "⬜", "Infrastructure detail"),
    ("support_email",          "⬜", "Infrastructure detail"),
])

add_table(doc, "SLA / timing (from stats object)", [
    ("stats.first_responded_at",     "✅", "Core SLA metric"),
    ("stats.resolved_at",            "✅", "Core SLA metric"),
    ("stats.closed_at",              "✅", ""),
    ("stats.status_updated_at",      "✅", "Last state change"),
    ("stats.reopened_at",            "✅", "Quality metric"),
    ("stats.agent_responded_at",     "✅", ""),
    ("stats.requester_responded_at", "✅", ""),
    ("stats.pending_since",          "⬜", "Low value for reporting"),
])

add_table(doc, "Escalation flags", [
    ("is_escalated", "✅", ""),
    ("fr_escalated", "✅", "First response escalated"),
    ("nr_escalated", "⬜", "Always false in sample"),
    ("nr_due_by",    "⬜", "Always null in sample"),
])

add_table(doc, "Content fields", [
    ("description",      "⬜", "HTML — ~35 MB extra per snapshot. Already in backfill file."),
    ("description_text", "⬜", "Plain text version — same concern."),
])

add_table(doc, "Email thread fields (85/200 or fewer non-null)", [
    ("cc_emails",        "⬜", "Low reporting value"),
    ("reply_cc_emails",  "⬜", "Low reporting value"),
    ("fwd_emails",       "⬜", "Low reporting value"),
    ("ticket_cc_emails", "⬜", "Low reporting value"),
    ("to_emails",        "⬜", "Low reporting value"),
])

add_table(doc, "Custom fields", [
    ("custom_fields.module",                       "✅", "Ticket category"),
    ("custom_fields.produkt",                      "✅", "Product tag"),
    ("custom_fields.error_id_type",                "⬜", "Mostly null"),
    ("custom_fields.cf_sub_type",                  "⬜", "Mostly null"),
    ("custom_fields.cf_followup",                  "⬜", "Mostly null"),
    ("custom_fields.cf_fsm_contact_name",          "⬜", "FSM / field service — always null"),
    ("custom_fields.cf_fsm_phone_number",          "⬜", "FSM / field service — always null"),
    ("custom_fields.cf_fsm_service_location",      "⬜", "FSM / field service — always null"),
    ("custom_fields.cf_fsm_appointment_start_time","⬜", "FSM / field service — always null"),
    ("custom_fields.cf_fsm_appointment_end_time",  "⬜", "FSM / field service — always null"),
])

doc.add_heading("Always null / zero-value — safe to drop entirely", level=2)
doc.add_paragraph(
    "ticket_bcc_emails, associated_tickets_count, source_info, structured_description, "
    "nr_due_by, spam, association_type, sentiment_score, initial_sentiment_score, "
    "internal_agent_id, internal_group_id, form_id"
)

doc.add_heading("Proposed field set summary", level=2)
doc.add_paragraph(
    "23 fields covering the full ticket lifecycle for SLA, team, and product reporting. "
    "description / description_text are excluded from nightly snapshots to keep file size "
    "manageable (~8-10 MB vs 44 MB); full-text content is already in the backfill file. "
    "Mark up this document with your changes and return for implementation."
)

doc.save("freshdesk_field_review.docx")
print("Saved: freshdesk_field_review.docx")
